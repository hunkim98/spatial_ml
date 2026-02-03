"""
Ordinance Parser - Extract text from municode docx files.

Parses the hierarchical docx files downloaded from municode and extracts
text content, preserving section structure.
"""

import re
from pathlib import Path
from dataclasses import dataclass, field
from typing import Iterator

from docx import Document
from docx.opc.exceptions import PackageNotFoundError


@dataclass
class Section:
    """A section of ordinance text."""
    title: str
    content: str
    level: int = 0  # Nesting level (0 = top level)


@dataclass 
class OrdinanceDocument:
    """A parsed ordinance document."""
    city: str
    state: str
    filename: str
    filepath: Path
    sections: list[Section] = field(default_factory=list)
    raw_text: str = ""
    
    @property
    def is_zoning_related(self) -> bool:
        """Check if this document is likely about zoning districts."""
        zoning_keywords = [
            'zoning', 'district', 'land use', 'residential', 'commercial',
            'industrial', 'agricultural', 'mixed use', 'overlay'
        ]
        text_lower = (self.filename + " " + self.raw_text[:2000]).lower()
        return any(kw in text_lower for kw in zoning_keywords)


class OrdinanceParser:
    """Parse ordinance documents from municode docx files."""
    
    # Municode uses this separator in filenames for hierarchy
    HIERARCHY_SEPARATOR = "⫸"
    
    def __init__(self, municode_path: Path):
        """
        Initialize parser for a city's ordinance documents.
        
        Args:
            municode_path: Path to city's municode folder (e.g., collector/tmp/zoning_ordinance/al/birmingham/)
        """
        self.municode_path = Path(municode_path)
        self.city = self.municode_path.name
        self.state = self.municode_path.parent.name
    
    def list_documents(self) -> list[Path]:
        """List all docx files in the municode folder."""
        return sorted(self.municode_path.glob("*.docx"))
    
    def list_zoning_documents(self) -> list[Path]:
        """List docx files that are likely zoning-related based on filename."""
        zoning_patterns = [
            r'zoning', r'district', r'land.?use', r'appendix.*[a-z].*zoning',
            r'title.*zoning', r'chapter.*zoning', r'article.*zoning'
        ]
        
        docs = []
        for doc_path in self.list_documents():
            filename_lower = doc_path.stem.lower()
            if any(re.search(pattern, filename_lower) for pattern in zoning_patterns):
                docs.append(doc_path)
        
        return docs
    
    def parse_document(self, doc_path: Path) -> OrdinanceDocument:
        """
        Parse a single docx file into an OrdinanceDocument.
        
        Args:
            doc_path: Path to the docx file
            
        Returns:
            OrdinanceDocument with extracted text and sections
        """
        try:
            doc = Document(doc_path)
        except PackageNotFoundError:
            # File might be corrupted or empty
            return OrdinanceDocument(
                city=self.city,
                state=self.state,
                filename=doc_path.name,
                filepath=doc_path,
                sections=[],
                raw_text=""
            )
        
        # Extract hierarchy from filename
        hierarchy = self._parse_filename_hierarchy(doc_path.stem)
        
        # Extract paragraphs
        paragraphs = []
        sections = []
        current_section_title = hierarchy[-1] if hierarchy else doc_path.stem
        current_section_content = []
        
        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue
            
            paragraphs.append(text)
            
            # Detect section headers (usually bold or specific patterns)
            is_header = self._is_section_header(para, text)
            
            if is_header and current_section_content:
                # Save previous section
                sections.append(Section(
                    title=current_section_title,
                    content="\n".join(current_section_content),
                    level=len(hierarchy)
                ))
                current_section_title = text
                current_section_content = []
            else:
                current_section_content.append(text)
        
        # Save last section
        if current_section_content:
            sections.append(Section(
                title=current_section_title,
                content="\n".join(current_section_content),
                level=len(hierarchy)
            ))
        
        raw_text = "\n\n".join(paragraphs)
        
        return OrdinanceDocument(
            city=self.city,
            state=self.state,
            filename=doc_path.name,
            filepath=doc_path,
            sections=sections,
            raw_text=raw_text
        )
    
    def _parse_filename_hierarchy(self, filename: str) -> list[str]:
        """
        Parse the hierarchy from a municode filename.
        
        Example:
            "Title 7 - BUILDING⫸CHAPTER 3. - PLANNING" -> ["Title 7 - BUILDING", "CHAPTER 3. - PLANNING"]
        """
        return filename.split(self.HIERARCHY_SEPARATOR)
    
    def _is_section_header(self, para, text: str) -> bool:
        """
        Determine if a paragraph is a section header.
        
        Heuristics:
        - Bold text
        - Starts with "Sec.", "Section", "Article", etc.
        - Short length
        - All caps
        """
        # Check for section patterns
        section_patterns = [
            r'^Sec\.\s*\d',
            r'^Section\s*\d',
            r'^Article\s+[IVX\d]',
            r'^ARTICLE\s+[IVX\d]',
            r'^Chapter\s+\d',
            r'^CHAPTER\s+\d',
            r'^\d+[\.-]\d+',  # e.g., "5.1" or "5-1"
        ]
        
        if any(re.match(pattern, text) for pattern in section_patterns):
            return True
        
        # Check if text is all caps and short (likely a header)
        if text.isupper() and len(text) < 100 and len(text.split()) < 15:
            return True
        
        # Check if paragraph has bold formatting
        try:
            runs = para.runs
            if runs and all(run.bold for run in runs if run.text.strip()):
                return len(text) < 200  # Headers are usually short
        except:
            pass
        
        return False
    
    def parse_all(self, zoning_only: bool = True) -> Iterator[OrdinanceDocument]:
        """
        Parse all documents in the municode folder.
        
        Args:
            zoning_only: If True, only parse documents with zoning-related filenames
            
        Yields:
            OrdinanceDocument for each parsed file
        """
        if zoning_only:
            doc_paths = self.list_zoning_documents()
        else:
            doc_paths = self.list_documents()
        
        for doc_path in doc_paths:
            yield self.parse_document(doc_path)
    
    def get_combined_zoning_text(self) -> str:
        """
        Get all zoning-related text combined into a single string.
        
        Useful for text search across the entire zoning ordinance.
        """
        texts = []
        for doc in self.parse_all(zoning_only=True):
            if doc.raw_text:
                texts.append(f"=== {doc.filename} ===\n{doc.raw_text}")
        
        # If no zoning-specific documents found, try all documents
        if not texts:
            for doc in self.parse_all(zoning_only=False):
                if doc.is_zoning_related and doc.raw_text:
                    texts.append(f"=== {doc.filename} ===\n{doc.raw_text}")
        
        return "\n\n".join(texts)


def find_district_sections(doc: OrdinanceDocument) -> list[Section]:
    """
    Find sections that define zoning districts.
    
    Looks for sections with titles like:
    - "R-1 Single Family District"
    - "Sec. 5.1. - R-1 district"
    - "ARTICLE III. - RESIDENTIAL DISTRICTS"
    """
    district_pattern = re.compile(
        r'(?:^|[\s\-\.])([A-Z]{1,3}[\-\s]?\d{1,2}[A-Z]?)(?:[\s\-\.]|$)',
        re.IGNORECASE
    )
    
    district_sections = []
    for section in doc.sections:
        # Check title for district code patterns
        if district_pattern.search(section.title):
            district_sections.append(section)
        # Also check content for "purpose" statements
        elif 'purpose of' in section.content.lower() and 'district' in section.content.lower():
            district_sections.append(section)
    
    return district_sections


if __name__ == "__main__":
    import sys
    
    # Example usage
    if len(sys.argv) > 1:
        municode_path = Path(sys.argv[1])
    else:
        # Default test path
        municode_path = Path("collector/tmp/zoning_ordinance/al/birmingham")
    
    if not municode_path.exists():
        print(f"Path not found: {municode_path}")
        sys.exit(1)
    
    parser = OrdinanceParser(municode_path)
    
    print(f"Parsing ordinances for {parser.city}, {parser.state.upper()}")
    print("=" * 60)
    
    # List all documents
    all_docs = parser.list_documents()
    zoning_docs = parser.list_zoning_documents()
    
    print(f"Total documents: {len(all_docs)}")
    print(f"Zoning-related: {len(zoning_docs)}")
    print()
    
    # Parse zoning documents
    print("Zoning documents:")
    for doc in parser.parse_all(zoning_only=True):
        print(f"  - {doc.filename}")
        print(f"    Sections: {len(doc.sections)}")
        print(f"    Text length: {len(doc.raw_text)} chars")
        
        # Find district sections
        district_sections = find_district_sections(doc)
        if district_sections:
            print(f"    District sections: {len(district_sections)}")
            for section in district_sections[:3]:
                print(f"      * {section.title[:60]}...")
